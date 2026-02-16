"""
Module A: Text Extraction Engine
=================================
Multi-path extraction pipeline:
1. Digital Document Extraction (PyPDF2 for PDF, python-docx for DOCX)
2. OCR Fallback using Tesseract + Poppler (triggered only when digital extraction fails)
3. Text Normalization (special chars, whitespace, lowercase)
4. Section Identification (Education, Experience, Skills, Certifications)

Output: Structured JSON with section-level text + full_text
"""

import os
import re
import json
import logging
import unicodedata
from pathlib import Path
from typing import Dict, Optional

logger = logging.getLogger(__name__)


class TextExtractionEngine:
    """
    Extracts and structures text from PDF or DOCX resume files.
    Falls back to OCR if digital extraction yields insufficient content.
    """

    # Minimum character count to consider digital extraction successful
    MIN_CHAR_THRESHOLD = 100

    def __init__(self):
        # Import section keywords from config
        from config.settings import SECTION_KEYWORDS, OCR_DPI, OCR_LANG
        self.section_keywords = SECTION_KEYWORDS
        self.ocr_dpi = OCR_DPI
        self.ocr_lang = OCR_LANG

    # ─── Public Interface ─────────────────────────────────────────────────────

    def extract(self, file_path: str) -> Dict:
        """
        Main entry point. Extracts, normalizes, and sections a resume file.

        Args:
            file_path: Absolute or relative path to PDF or DOCX file

        Returns:
            Structured dict with education, experience, skills, certifications, full_text
        """
        file_path = str(Path(file_path).resolve())
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = Path(file_path).suffix.lower()
        logger.info(f"[TextExtraction] Processing: {file_path} ({ext})")

        # Step 1: Digital extraction
        raw_text = ""
        if ext == ".pdf":
            raw_text = self._extract_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            raw_text = self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        # Step 2: OCR fallback if digital extraction insufficient
        if len(raw_text.strip()) < self.MIN_CHAR_THRESHOLD:
            logger.warning(
                f"[TextExtraction] Digital extraction insufficient "
                f"({len(raw_text)} chars). Activating OCR fallback."
            )
            raw_text = self._ocr_fallback(file_path, ext)

        # Step 3: Normalize text
        normalized = self._normalize_text(raw_text)

        # Step 4: Identify sections
        sections = self._identify_sections(normalized)
        sections["full_text"] = normalized

        logger.info(
            f"[TextExtraction] Complete. Sections found: "
            f"{[k for k, v in sections.items() if v and k != 'full_text']}"
        )
        return sections

    # ─── 1. Digital PDF Extraction ────────────────────────────────────────────

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from a digitally-created PDF using pypdf (or PyPDF2)."""
        text_parts = []
        # Try pypdf first (modern), fall back to PyPDF2
        try:
            import pypdf
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text() or ""
                        text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"[PDF] Failed to extract page {page_num}: {e}")
            full_text = "\n".join(text_parts)
            logger.info(f"[PDF] Extracted {len(full_text)} chars via pypdf.")
            return full_text
        except ImportError:
            pass

        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text() or ""
                        text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"[PDF] Failed to extract page {page_num}: {e}")
            full_text = "\n".join(text_parts)
            logger.info(f"[PDF] Extracted {len(full_text)} chars via PyPDF2.")
            return full_text
        except ImportError:
            logger.error("[PDF] Neither pypdf nor PyPDF2 is installed.")
            return ""
        except Exception as e:
            logger.error(f"[PDF] Extraction failed: {e}")
            return ""

    # ─── 2. DOCX Extraction ───────────────────────────────────────────────────

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file using python-docx."""
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract table cells (skills tables in resumes)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            full_text = "\n".join(text_parts)
            logger.info(f"[DOCX] Extracted {len(full_text)} chars, {len(doc.paragraphs)} paragraphs.")
            return full_text
        except ImportError:
            logger.error("[DOCX] python-docx not installed. Run: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"[DOCX] Extraction failed: {e}")
            return ""

    # ─── 3. OCR Fallback ──────────────────────────────────────────────────────

    def _ocr_fallback(self, file_path: str, ext: str) -> str:
        """
        Fallback OCR pipeline using Tesseract + Poppler.
        Note: Accuracy may reduce depending on scan quality.
        """
        try:
            import pytesseract
            from PIL import Image

            images = []

            if ext == ".pdf":
                # Use pdf2image (Poppler backend) to convert PDF pages to images
                try:
                    from pdf2image import convert_from_path
                    logger.info(f"[OCR] Converting PDF to images (DPI={self.ocr_dpi})...")
                    images = convert_from_path(file_path, dpi=self.ocr_dpi)
                except ImportError:
                    logger.error("[OCR] pdf2image not installed. Run: pip install pdf2image")
                    return ""
            else:
                # For non-PDF files, try direct OCR if image format
                try:
                    img = Image.open(file_path)
                    images = [img]
                except Exception:
                    logger.error(f"[OCR] Cannot open {ext} as image.")
                    return ""

            if not images:
                logger.error("[OCR] No images generated for OCR.")
                return ""

            text_parts = []
            for i, image in enumerate(images):
                try:
                    page_text = pytesseract.image_to_string(
                        image,
                        lang=self.ocr_lang,
                        config="--psm 6"  # Assume uniform block of text
                    )
                    text_parts.append(page_text)
                    logger.debug(f"[OCR] Page {i+1}: {len(page_text)} chars")
                except Exception as e:
                    logger.warning(f"[OCR] Failed on page {i+1}: {e}")

            full_text = "\n".join(text_parts)
            logger.info(f"[OCR] Extracted {len(full_text)} chars from {len(images)} pages.")
            return full_text

        except ImportError:
            logger.error("[OCR] pytesseract not installed. Run: pip install pytesseract")
            return ""
        except Exception as e:
            logger.error(f"[OCR] Unexpected error: {e}")
            return ""

    # ─── 4. Text Normalization ────────────────────────────────────────────────

    def _normalize_text(self, text: str) -> str:
        """
        Normalize extracted text for downstream processing.
        Steps: unicode normalization, special chars, whitespace, lowercase.
        """
        if not text:
            return ""

        # Normalize unicode characters (smart quotes, dashes, etc.)
        text = unicodedata.normalize("NFKD", text)

        # Replace common formatting artifacts
        text = text.replace("\x00", "")          # Null bytes
        text = text.replace("\r\n", "\n")         # Windows line endings
        text = text.replace("\r", "\n")           # Mac line endings
        text = text.replace("\t", " ")            # Tabs to spaces
        text = text.replace("\u2022", "- ")       # Bullet points
        text = text.replace("\u2019", "'")        # Smart apostrophe
        text = text.replace("\u201c", '"')        # Smart quote open
        text = text.replace("\u201d", '"')        # Smart quote close

        # Remove non-ASCII characters except common punctuation
        text = re.sub(r"[^\x00-\x7F]+", " ", text)

        # Remove special characters that aren't content-bearing
        # Keep: alphanumeric, spaces, hyphens, slashes, commas, periods, parentheses
        text = re.sub(r"[^a-zA-Z0-9\s\-\/\,\.\_\(\)\@\+\#]", " ", text)

        # Normalize whitespace: collapse multiple spaces
        text = re.sub(r" {2,}", " ", text)

        # Remove redundant newlines (keep double newlines as paragraph breaks)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Strip leading/trailing whitespace per line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(line for line in lines if line)

        # Lowercase standardization
        text = text.lower()

        return text.strip()

    # ─── 5. Section Identification ────────────────────────────────────────────

    def _identify_sections(self, normalized_text: str) -> Dict[str, str]:
        """
        Identify resume sections using rule-based + keyword-based stratified detection.

        Strategy:
        1. Split text into lines
        2. Detect section header lines (short, keyword-matching lines)
        3. Assign subsequent lines to detected sections
        4. Fall back to full-text keyword search if header detection fails

        Returns:
            Dict with keys: education, experience, skills, certifications
        """
        sections = {
            "education": "",
            "experience": "",
            "skills": "",
            "certifications": ""
        }

        lines = normalized_text.split("\n")
        current_section = None
        section_content = {k: [] for k in sections}

        def _is_section_header(line: str) -> Optional[str]:
            """Determine if a line is a section header and return section name."""
            line_clean = line.strip().lower()
            # Section headers are typically short (< 40 chars) and keyword-matching
            if len(line_clean) > 60:
                return None
            for section_name, keywords in self.section_keywords.items():
                for keyword in keywords:
                    if keyword in line_clean:
                        return section_name
            return None

        # Pass 1: Header-based section detection
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue

            detected_section = _is_section_header(line_stripped)
            if detected_section:
                current_section = detected_section
                continue  # Don't add the header itself to content

            if current_section:
                section_content[current_section].append(line_stripped)

        # Build section strings from gathered content
        for section_name in sections:
            if section_content[section_name]:
                sections[section_name] = "\n".join(section_content[section_name])

        # Pass 2: Keyword search fallback for any empty sections
        for section_name, keywords in self.section_keywords.items():
            if sections[section_name]:
                continue  # Already have content

            # Search for keyword context in full text
            for keyword in keywords[:3]:  # Check top 3 keywords
                pattern = rf"(?i){re.escape(keyword)}[\s\S]{{0,500}}"
                match = re.search(pattern, normalized_text)
                if match:
                    context = match.group(0)[:500]
                    sections[section_name] = context
                    break

        return sections

    # ─── Utility ─────────────────────────────────────────────────────────────

    def extract_from_text(self, raw_text: str) -> Dict:
        """
        Process pre-loaded text string directly (useful for testing).

        Args:
            raw_text: Raw resume text

        Returns:
            Structured sections dict
        """
        normalized = self._normalize_text(raw_text)
        sections = self._identify_sections(normalized)
        sections["full_text"] = normalized
        return sections

    def to_json(self, result: Dict, indent: int = 2) -> str:
        """Serialize extraction result to JSON string."""
        return json.dumps(result, indent=indent, ensure_ascii=False)
